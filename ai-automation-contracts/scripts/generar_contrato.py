#!/usr/bin/env python3
"""
Genera un contrato .docx a partir de una plantilla markdown y un JSON de datos.

Uso:
    python3 generar_contrato.py <plantilla.md> <datos.json> <salida.docx>

El JSON de datos es un diccionario plano: {"CLIENTE_CUIT": "30-...", ...}
Las claves corresponden a los {{PLACEHOLDERS}} de la plantilla.
Cualquier placeholder sin valor se reemplaza por [ ... ] resaltado para completar a mano.
"""

import json
import os
import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor


PLACEHOLDER = re.compile(r"\{\{([A-Z0-9_]+)\}\}")
BOLD = re.compile(r"\*\*(.+?)\*\*")

LOGO_PATH = os.path.join(os.path.dirname(__file__), "..", "assets", "logo.png")
MARCA_NOMBRE = "AURELIO AGENCY"
MARCA_WEB = "aurelioagency.com"


def render_template(text, datos):
    """Sustituye {{CLAVE}} por su valor. Los faltantes quedan marcados."""
    faltantes = []

    def repl(m):
        clave = m.group(1)
        valor = datos.get(clave)
        if valor is None or str(valor).strip() == "":
            faltantes.append(clave)
            return f"[COMPLETAR: {clave}]"
        return str(valor)

    return PLACEHOLDER.sub(repl, text), sorted(set(faltantes))


def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size, color in (
        ("Heading 1", 15, RGBColor(0, 0, 0)),
        ("Heading 2", 12, RGBColor(0, 0, 0)),
        ("Heading 3", 11, RGBColor(0, 0, 0)),
    ):
        st = doc.styles[name]
        st.font.name = "Arial"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(14)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def add_page_number(paragraph):
    """Inserta el campo de número de página de Word en el párrafo dado."""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def setup_header_footer(doc):
    """Encabezado con logo + nombre de marca y pie con sitio web + n° de página."""
    section = doc.sections[0]
    right_edge = section.page_width - section.left_margin - section.right_margin

    header_par = section.header.paragraphs[0]
    header_par.paragraph_format.tab_stops.add_tab_stop(right_edge, WD_TAB_ALIGNMENT.RIGHT)
    if os.path.exists(LOGO_PATH):
        header_par.add_run().add_picture(LOGO_PATH, height=Cm(0.45))
    marca_run = header_par.add_run(f"\t{MARCA_NOMBRE}")
    marca_run.bold = True
    marca_run.font.size = Pt(8)
    marca_run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    footer_par = section.footer.paragraphs[0]
    footer_par.paragraph_format.tab_stops.add_tab_stop(right_edge, WD_TAB_ALIGNMENT.RIGHT)
    web_run = footer_par.add_run(MARCA_WEB)
    web_run.font.size = Pt(8)
    web_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    footer_par.add_run("\t")
    add_page_number(footer_par)


def add_portada(doc):
    """Bloque de portada con logo y marca antes del título del documento."""
    if os.path.exists(LOGO_PATH):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.add_run().add_picture(LOGO_PATH, height=Cm(1.8))

    p_marca = doc.add_paragraph()
    p_marca.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_marca = p_marca.add_run(MARCA_NOMBRE)
    r_marca.bold = True
    r_marca.font.size = Pt(12)

    p_web = doc.add_paragraph()
    p_web.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_web = p_web.add_run(MARCA_WEB)
    r_web.font.size = Pt(8.5)
    r_web.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    p_web.paragraph_format.space_after = Pt(20)


def add_runs(par, text):
    """Escribe texto en un párrafo respetando **negrita** y marcas [COMPLETAR: X]."""
    pos = 0
    for m in BOLD.finditer(text):
        if m.start() > pos:
            add_plain(par, text[pos:m.start()], bold=False)
        add_plain(par, m.group(1), bold=True)
        pos = m.end()
    if pos < len(text):
        add_plain(par, text[pos:], bold=False)


def add_plain(par, text, bold):
    """Divide el texto en tramos normales y marcas de completar (resaltadas en rojo)."""
    for chunk in re.split(r"(\[COMPLETAR: [A-Z0-9_]+\])", text):
        if not chunk:
            continue
        run = par.add_run(chunk)
        run.bold = bold
        if chunk.startswith("[COMPLETAR:"):
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
            run.bold = True


def parse_table(lines, i):
    """Devuelve (filas, siguiente_indice) para un bloque de tabla markdown."""
    filas = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        celdas = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{2,}:?", c) for c in celdas if c):
            filas.append(celdas)
        i += 1
    return filas, i


def build_docx(md, salida):
    doc = Document()
    setup_styles(doc)
    setup_header_footer(doc)
    add_portada(doc)

    lines = md.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Separador: solo genera salto de página antes de un ANEXO
        if stripped == "---":
            j = i + 1
            while j < len(lines) and not lines[j].strip():
                j += 1
            if j < len(lines) and re.match(r"^#\s+ANEXO", lines[j].strip()):
                p = doc.add_paragraph()
                p.add_run().add_break(WD_BREAK.PAGE)
            i += 1
            continue

        if stripped == "<br><br>":
            doc.add_paragraph()
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("|"):
            filas, i = parse_table(lines, i)
            if not filas:
                continue
            ncols = max(len(f) for f in filas)
            tbl = doc.add_table(rows=0, cols=ncols)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            for fila in filas:
                cells = tbl.add_row().cells
                for k in range(ncols):
                    texto = fila[k] if k < len(fila) else ""
                    par = cells[k].paragraphs[0]
                    par.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    add_runs(par, texto)
            doc.add_paragraph()
            continue

        m = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if m:
            nivel = len(m.group(1))
            par = doc.add_heading(level=nivel)
            add_runs(par, m.group(2))
            if nivel == 1:
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        if re.match(r"^[-*]\s+", stripped):
            par = doc.add_paragraph(style="List Bullet")
            add_runs(par, re.sub(r"^[-*]\s+", "", stripped))
            i += 1
            continue

        par = doc.add_paragraph()
        add_runs(par, stripped)
        i += 1

    doc.save(salida)


def main():
    if len(sys.argv) != 4:
        print(__doc__)
        sys.exit(1)

    plantilla, datos_path, salida = sys.argv[1:4]

    with open(plantilla, encoding="utf-8") as f:
        md = f.read()
    with open(datos_path, encoding="utf-8") as f:
        datos = json.load(f)

    md, faltantes = render_template(md, datos)
    build_docx(md, salida)

    print(f"Generado: {salida}")
    if faltantes:
        print("\nCampos sin completar (marcados en rojo en el documento):")
        for c in faltantes:
            print(f"  - {c}")
    else:
        print("Todos los campos fueron completados.")


if __name__ == "__main__":
    main()
